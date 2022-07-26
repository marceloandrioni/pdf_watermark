#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Add a watermark to a pdf file.

Author: Marcelo Andrioni
https://github.com/marceloandrioni

"""

import os
import sys
import shutil
import tempfile
import argparse
from pathlib import Path
import numpy as np
from docx import Document
from docx.shared import RGBColor, Mm, Pt
from gooey import Gooey, GooeyParser
from subprocess import Popen
from pikepdf import Pdf, Page, Encryption, Permissions

# the docx conversion to pdf can be done with docx2pdf/word (windows only) or
# using libreoffice (windows and linux)
HAS_LIBREOFFICE = True if shutil.which('libreoffice') else False
try:

    # find Microsoft Word
    import win32com.client
    _ = win32com.client.Dispatch("Word.Application")

    import docx2pdf

    HAS_DOCX2PDF = True

except:

    HAS_DOCX2PDF = False


# Use CLI (instead of GUI) if the CLI arguments were passed.
# https://github.com/chriskiehl/Gooey/issues/449#issuecomment-534056010
if len(sys.argv) > 1:
    if '--ignore-gooey' not in sys.argv:
        sys.argv.append('--ignore-gooey')


def create_wmark_docx(text):

    # create water mark file in temporary folder to overlay in main pdf file
    wmark_docx = (Path(tempfile.gettempdir())
                  / 'watermark_{:06d}.docx'.format(np.random.randint(1_000_000)))

    # create docx
    document = Document()

    # set size as A4
    # https://stackoverflow.com/a/54757281/9707202
    for section in document.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.header_distance = Mm(12.7)

        # set a very low footer
        section.footer_distance = Mm(5)   # Mm(12.7)

    # p = document.add_paragraph()
    p = document.sections[0].footer.add_paragraph()   # add the text as footer

    # 0 left, 1 center, 2 right, 3 justify
    p.alignment = 1

    run = p.add_run(text)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(255, 0, 0)

    document.save(wmark_docx)

    return wmark_docx


def libreoffice_docx2pdf(input_docx, out_folder):
    """Convert docx file to pdf using libreoffice.
    Reference: https://stackoverflow.com/a/56067358/9707202
    Author: https://stackoverflow.com/users/7037499/dfresh22
    """

    # get libreoffice executable location
    libre_office = shutil.which('libreoffice')
    if not libre_office:
        raise ValueError('Could not find libreoffice executable location in path.')

    p = Popen([libre_office, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    # print([libre_office, '--convert-to', 'pdf', input_docx])
    p.communicate()


def convert_docx_to_pdf(wmark_docx):

    wmark_pdf = wmark_docx.with_suffix('.pdf')

    # convert docx to pdf
    if HAS_LIBREOFFICE:
        libreoffice_docx2pdf(wmark_docx, wmark_pdf.parent)
    else:
        docx2pdf.convert(wmark_docx, wmark_pdf)

    return wmark_pdf


def user_args():

    description = 'Add a watermark to a pdf file.'
    parser = GooeyParser(description=description,
                         allow_abbrev=False)

    parser.add_argument('infile',
                        type=lambda x: Path(x),
                        help='Input pdf file.',
                        widget='FileChooser',
                        gooey_options={
                            'wildcard': 'PDF file (*.pdf)|*.pdf',
                            'message': 'Select input pdf file'})

    parser.add_argument('outfile',
                        type=lambda x: Path(x),
                        help='Output pdf file.',
                        widget='FileSaver',
                        gooey_options={
                            'wildcard': 'PDF file (*.pdf)|*.pdf',
                            'message': 'Select output pdf file'})

    # hack to set the name of the group in Gooey
    # https://stackoverflow.com/questions/53498352/is-there-a-solution-for-required-mutually-exclusive-arguments-listed-as-optional
    group_wmark = parser.add_argument_group(title='Watermark options')

    group_wmark2 = group_wmark.add_mutually_exclusive_group(required=True)

    group_wmark2.add_argument('-wt', '--watermark_text',
                              dest='wmark_text',
                              help='Text for watermark.',
                              default='')

    group_wmark2.add_argument('-wd', '--watermark_docx',
                              type=lambda x: Path(x),
                              dest='wmark_docx',
                              help='Watermark docx file.',
                              widget='FileChooser',
                              gooey_options={
                                  'wildcard': 'Docx file (*.docx)|*.docx',
                                  'message': 'Select watermark docx file'})

    group_wmark2.add_argument('-wp', '--watermark_pdf',
                              type=lambda x: Path(x),
                              dest='wmark_pdf',
                              help='Watermark pdf file.',
                              widget='FileChooser',
                              gooey_options={
                                  'wildcard': 'PDF file (*.pdf)|*.pdf',
                                  'message': 'Select watermark pdf file'})

    args = parser.parse_args()

    if not args.wmark_pdf and not any((HAS_LIBREOFFICE, HAS_DOCX2PDF)):
        raise ValueError('libreoffice or docx2pdf/word must be installed to '
                         'create a pdf file from the watermark text.')

    # additional checks in case user bypass Gooey to use argparse directly
    if [args.infile.suffix, args.outfile.suffix] != ['.pdf', '.pdf']:
        raise argparse.ArgumentTypeError('Input/Output files must be pdf files.')

    if not args.infile.exists():
        raise argparse.ArgumentTypeError(f"Input file '{args.infile}' does "
                                         "not exist.")

    if args.outfile.exists() and os.path.samefile(args.infile, args.outfile):
        raise argparse.ArgumentTypeError("Input/Output files can't be the same.")

    if args.wmark_docx and args.wmark_docx.suffix != '.docx':
        raise argparse.ArgumentTypeError('Watermark docx file must have .docx '
                                         'extension.')

    if args.wmark_pdf and args.wmark_pdf.suffix != '.pdf':
        raise argparse.ArgumentTypeError('Watermark pdf file must have .pdf '
                                         'extension.')

    return args


@Gooey(required_cols=1,
       default_size=(610, 800),
       progress_regex=r"^Page (?P<current>\d+)/(?P<total>\d+)$",
       progress_expr="current / total * 100")
def main():

    args = user_args()

    # if wmark_pdf was given, just use it
    # if wmark_docx was given, convert to pdf and use it
    # if wmark_text was given, create docx, convert to pdf and use it
    if args.wmark_pdf:
        wmark_pdf = args.wmark_pdf
    elif args.wmark_docx:
        wmark_pdf = convert_docx_to_pdf(args.wmark_docx)
    else:
        wmark_docx = create_wmark_docx(args.wmark_text)
        wmark_pdf = convert_docx_to_pdf(wmark_docx)

    wmark = Pdf.open(wmark_pdf)
    thumbnail = Page(wmark.pages[0])

    print(f'Input file: {args.infile}')

    pdf = Pdf.open(args.infile)

    # merge water mark file with each page of input pdf file
    for page in pdf.pages:

        print(f'Page {page.index + 1}/{len(pdf.pages)}')

        page.add_overlay(thumbnail)

    # Do not allow a regular user to modify the file.
    # This way a user can't simply remove the watermark using a pdf editor like
    # LibreOffice Draw.
    allow = Permissions(accessibility=True,
                        extract=True,
                        modify_annotation=False,
                        modify_assembly=False,
                        modify_form=False,
                        modify_other=False,
                        print_lowres=True,
                        print_highres=True)
    encryption = Encryption(user='', owner='admin123', allow=allow)

    print(f'Output file: {args.outfile}')
    pdf.save(args.outfile, linearize=True, encryption=encryption)

    pdf.close()

    wmark.close()

    # remove temporary files (keep template docx/pdf file given by user)
    if args.wmark_text:
        os.remove(wmark_docx)
        os.remove(wmark_pdf)
    elif args.wmark_docx:
        os.remove(wmark_pdf)

    print('Done!')


if __name__ == '__main__':
    main()
